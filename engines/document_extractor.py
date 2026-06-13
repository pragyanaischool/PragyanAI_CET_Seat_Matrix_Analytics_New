import io
import os
import tempfile
import pandas as pd
from docx import Document

# ==========================================================
# Optional Extraction Engines
# ==========================================================

try:
    from docling.document_converter import (
        DocumentConverter as IBMDoclingConverter
    )
except ImportError:
    IBMDoclingConverter = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None


class PragyanDocumentExtractor:
    """
    Enterprise-grade document extraction engine.

    Extraction Priority:

    PDF
      -> IBM Docling (Primary)
      -> pdfplumber (Fallback)

    Excel
      -> Pandas

    CSV
      -> Pandas

    DOCX
      -> python-docx
    """

    # ======================================================
    # Public Entry Point
    # ======================================================

    @staticmethod
    def extract_to_text_stream(
        file_bytes: bytes,
        file_name: str
    ) -> str:

        if not file_name or "." not in file_name:
            raise ValueError(
                "Invalid document reference: Missing file extension."
            )

        ext = file_name.split(".")[-1].lower()

        if ext == "pdf":
            return PragyanDocumentExtractor._extract_pdf_pipeline(
                file_bytes
            )

        elif ext in ["xlsx", "xls"]:
            return PragyanDocumentExtractor._extract_spreadsheet(
                file_bytes
            )

        elif ext == "csv":
            return PragyanDocumentExtractor._extract_csv(
                file_bytes
            )

        elif ext in ["docx", "doc"]:
            return PragyanDocumentExtractor._extract_docx(
                file_bytes
            )

        raise ValueError(
            f"Unsupported file type: .{ext}"
        )

    # ======================================================
    # PDF Extraction
    # ======================================================

    @staticmethod
    def _extract_pdf_pipeline(
        file_bytes: bytes
    ) -> str:

        absolute_temp_path = None

        with tempfile.NamedTemporaryFile(
            delete=False,
            suffix=".pdf"
        ) as temp_pdf:

            temp_pdf.write(file_bytes)
            absolute_temp_path = os.path.abspath(
                temp_pdf.name
            )

        try:

            # --------------------------------------------------
            # Strategy A: IBM Docling
            # --------------------------------------------------

            if IBMDoclingConverter is not None:

                try:

                    print(
                        "[Docling] Starting document conversion..."
                    )

                    converter = IBMDoclingConverter()

                    result = converter.convert(
                        absolute_temp_path
                    )

                    markdown_output = (
                        result.document.export_to_markdown()
                    )

                    if markdown_output:

                        markdown_output = (
                            markdown_output
                            .replace("\u00a0", " ")
                            .replace("\ufeff", "")
                        )

                        markdown_output = (
                            markdown_output.strip()
                        )

                    print(
                        f"[Docling] Extracted "
                        f"{len(markdown_output)} characters"
                    )

                    if markdown_output:
                        print(
                            "[Docling] Extraction successful."
                        )
                        return markdown_output

                except Exception as e:

                    print(
                        "[Extractor Alert] "
                        f"Docling failed: {str(e)}"
                    )

            # --------------------------------------------------
            # Strategy B: pdfplumber fallback
            # --------------------------------------------------

            if pdfplumber is not None:

                print(
                    "[Fallback] Running pdfplumber extraction..."
                )

                text_accum = []

                with pdfplumber.open(
                    absolute_temp_path
                ) as pdf:

                    for idx, page in enumerate(
                        pdf.pages,
                        start=1
                    ):

                        try:

                            page_text = page.extract_text(
                                layout=True
                            )

                            if page_text:

                                page_text = (
                                    page_text
                                    .replace("\u00a0", " ")
                                    .strip()
                                )

                                text_accum.append(
                                    f"\n--- PAGE {idx} ---\n"
                                    f"{page_text}"
                                )

                        except Exception as page_error:

                            print(
                                f"[Fallback Warning] "
                                f"Page {idx}: "
                                f"{str(page_error)}"
                            )

                final_text = "\n".join(text_accum)

                print(
                    f"[Fallback] Extracted "
                    f"{len(final_text)} characters"
                )

                return final_text

            raise RuntimeError(
                "No PDF extraction engine available."
            )

        finally:

            if (
                absolute_temp_path
                and os.path.exists(
                    absolute_temp_path
                )
            ):

                try:
                    os.remove(
                        absolute_temp_path
                    )
                except Exception:
                    pass

    # ======================================================
    # DOCX Extraction
    # ======================================================

    @staticmethod
    def _extract_docx(
        file_bytes: bytes
    ) -> str:

        try:

            doc = Document(
                io.BytesIO(file_bytes)
            )

            text_accum = []

            # Paragraphs
            for paragraph in doc.paragraphs:

                text = paragraph.text.strip()

                if text:
                    text_accum.append(text)

            # Tables
            for table in doc.tables:

                for row in table.rows:

                    cells = [
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    ]

                    if cells:
                        text_accum.append(
                            " | ".join(cells)
                        )

            return "\n".join(text_accum)

        except Exception as e:

            raise RuntimeError(
                f"DOCX extraction failed: {str(e)}"
            )

    # ======================================================
    # Excel Extraction
    # ======================================================

    @staticmethod
    def _extract_spreadsheet(
        file_bytes: bytes
    ) -> str:

        try:

            excel_file = pd.ExcelFile(
                io.BytesIO(file_bytes)
            )

            text_accum = []

            for sheet_name in excel_file.sheet_names:

                try:

                    df = (
                        excel_file.parse(sheet_name)
                        .dropna(how="all")
                    )

                    if df.empty:
                        continue

                    df.columns = [
                        str(c).strip()
                        for c in df.columns
                    ]

                    sheet_text = df.to_csv(
                        index=False
                    )

                    text_accum.append(
                        f"\n--- SHEET: {sheet_name} ---\n"
                        f"{sheet_text}"
                    )

                except Exception as sheet_error:

                    print(
                        f"[Excel Warning] "
                        f"{sheet_name}: "
                        f"{str(sheet_error)}"
                    )

            return "\n".join(text_accum)

        except Exception as e:

            raise RuntimeError(
                f"Excel extraction failed: {str(e)}"
            )

    # ======================================================
    # CSV Extraction
    # ======================================================

    @staticmethod
    def _extract_csv(
        file_bytes: bytes
    ) -> str:

        try:

            df = pd.read_csv(
                io.BytesIO(file_bytes),
                encoding="utf-8",
                on_bad_lines="skip"
            )

            df = df.dropna(how="all")

            return df.to_csv(
                index=False
            )

        except UnicodeDecodeError:

            try:

                df = pd.read_csv(
                    io.BytesIO(file_bytes),
                    encoding="latin-1",
                    on_bad_lines="skip"
                )

                return df.to_csv(
                    index=False
                )

            except Exception as e:

                raise RuntimeError(
                    f"CSV extraction failed: {str(e)}"
                )

        except Exception as e:

            raise RuntimeError(
                f"CSV extraction failed: {str(e)}"
            )
